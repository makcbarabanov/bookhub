"""Book export helpers (HTML, DOCX, PDF)."""

from __future__ import annotations

import html as html_lib
import io
import re
from typing import Any

from bs4 import BeautifulSoup
from docx import Document
from fpdf import FPDF


def _plain_text(content_html: str) -> str:
    soup = BeautifulSoup(content_html or "", "html.parser")
    for br in soup.find_all("br"):
        br.replace_with("\n")
    for p in soup.find_all("p"):
        p.append("\n")
    text = soup.get_text("\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def build_html_export(book_title: str, chapters: list[dict[str, Any]]) -> str:
    parts = [
        "<!DOCTYPE html>",
        '<html lang="ru"><head><meta charset="UTF-8">',
        f"<title>{html_lib.escape(book_title)}</title>",
        "<style>body{font-family:Georgia,serif;max-width:720px;margin:2rem auto;line-height:1.6}",
        "h2{margin-top:2rem;color:#333}</style>",
        "</head><body>",
        f"<h1>{html_lib.escape(book_title)}</h1>",
    ]
    for i, ch in enumerate(chapters, start=1):
        parts.append(
            f'<section id="{html_lib.escape(ch["ch_id"])}">'
            f"<h2>{i}. {html_lib.escape(ch['title'])}</h2>"
            f'{ch.get("content") or ""}'
            "</section>"
        )
    parts.append("</body></html>")
    return "\n".join(parts)


def build_docx_export(book_title: str, chapters: list[dict[str, Any]]) -> bytes:
    doc = Document()
    doc.add_heading(book_title, 0)
    for i, ch in enumerate(chapters, start=1):
        doc.add_heading(f"{i}. {ch['title']}", level=1)
        text = _plain_text(ch.get("content") or "")
        if text:
            for block in text.split("\n\n"):
                block = block.strip()
                if block:
                    doc.add_paragraph(block)
        else:
            doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class _BookPDF(FPDF):
    def __init__(self) -> None:
        super().__init__()
        self.add_font("DejaVu", "", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
        self.add_font("DejaVu", "B", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
        self.set_auto_page_break(auto=True, margin=15)


def build_pdf_export(book_title: str, chapters: list[dict[str, Any]]) -> bytes:
    pdf = _BookPDF()
    pdf.add_page()
    pdf.set_font("DejaVu", "B", 16)
    pdf.multi_cell(0, 10, book_title)
    pdf.ln(4)
    for i, ch in enumerate(chapters, start=1):
        pdf.set_font("DejaVu", "B", 13)
        pdf.multi_cell(0, 8, f"{i}. {ch['title']}")
        pdf.ln(2)
        pdf.set_font("DejaVu", "", 11)
        text = _plain_text(ch.get("content") or "")
        if text:
            pdf.multi_cell(0, 6, text)
        pdf.ln(4)
    out = pdf.output()
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return out.encode("latin-1")
